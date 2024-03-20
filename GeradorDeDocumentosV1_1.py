import os
import base64
import requests
import tkinter as tk
from docx import Document
from ttkthemes import ThemedTk
from PIL import Image, ImageTk, ImageFilter
from docx.shared import Pt, Inches
from azure.devops.connection import Connection
from tkinter import ttk, messagebox, filedialog
from msrest.authentication import BasicAuthentication
import shutil


import json

def load_version():
    try:
        with open("config.json", "r") as config_file:
            config_data = json.load(config_file)
            return config_data["current_version"]
    except (FileNotFoundError, KeyError):
        # Se o arquivo de configuração não existir ou a chave current_version não estiver presente
        # retorna a versão padrão
        return "1.0"

def verificar_atualizacoes():
    try:
        # Carregar os dados do arquivo config.json
        config_data = load_version()

        # URL do seu repositório GitHub
        repo_url = "https://github.com/RenzoDTavares/AssistenteDeTestes"

        # Fazendo a requisição GET para obter informações sobre o repositório
        response = requests.get(repo_url)

        # Verificando se a requisição foi bem-sucedida
        if response.status_code == 200:
            repo_info = response.json()

            # Obtendo a última versão do repositório
            latest_version = repo_info['tag_name']

            # Obtendo a versão atual do aplicativo
            current_version = config_data

            # Comparando a última versão com a versão atual
            if latest_version > current_version:
                # Baixando e atualizando os arquivos do repositório
                download_url = repo_info['zipball_url']
                download_response = requests.get(download_url)
                with open("update.zip", "wb") as update_file:
                    update_file.write(download_response.content)

                # Extrair o arquivo zip com as atualizações
                import zipfile
                with zipfile.ZipFile("update.zip", "r") as zip_ref:
                    zip_ref.extractall("update")

                # Copiando os arquivos atualizados para o diretório atual
                for root, dirs, files in os.walk("update"):
                    for file in files:
                        shutil.copy(os.path.join(root, file), os.getcwd())

                # Limpar diretório de atualização
                shutil.rmtree("update")
                os.remove("update.zip")

                # Atualizar o número da versão no arquivo config.json
                with open("config.json", "w") as config_file:
                    config_data["current_version"] = latest_version
                    json.dump(config_data, config_file)

                # Mostrar mensagem de sucesso
                messagebox.showinfo("Atualização", "Atualização bem-sucedida. Por favor, reinicie o aplicativo.")

            else:
                messagebox.showinfo("Atualização", "Você já tem a versão mais recente do aplicativo.")

        else:
            messagebox.showerror("Erro", "Não foi possível verificar atualizações no momento.")

    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao verificar atualizações: {str(e)}")


# Função para iniciar a busca para o browser
def on_browse():
    entries[0].delete(0, tk.END)
    entries[0].insert(tk.END, filedialog.askopenfilename(filetypes=[("Word files", "*.docx")]))

# Função para iniciar a busca para o diretório
def on_select_image_dir():
    entries[1].delete(0, tk.END)
    entries[1].insert(tk.END, filedialog.askdirectory())

# Função que remove OOTB do Tkinker, que seleciona o item do dropdown
def desativar_selecao(event):
    ambiente_dropdown.selection_clear()
    project_dropdown.selection_clear()
    
# Função para cadastrar ou renovar a chave do Azure DevOps
def cadastrar_renovar_chave():
        chave_devops = tk.simpledialog.askstring("Chave do Azure DevOps", "Digite a sua chave do Azure DevOps:")
        if chave_devops:
            chave_codificada = base64.b64encode(chave_devops.encode()).decode()
            with open("chave_devops.txt", "w") as file:
                file.write(chave_codificada)
            messagebox.showinfo("Chave Cadastrada", "Chave do Azure DevOps cadastrada com sucesso.")
       
# Função para decodificar a chave do Azure DevOps
def decodificar_chave():
    try:
        with open("chave_devops.txt", "r") as file:
            chave_codificada = file.read().strip()
            chave_decodificada = base64.b64decode(chave_codificada.encode()).decode()
            return chave_decodificada
    except FileNotFoundError:        
        return False

# Função que valida se as credenciais do DevOps retornam ao menos um e-mail @br.ey para validar se a chave esta no mínimo OK
def validar_credenciais(organization_url, personal_access_token):
    try:
        url = f"{organization_url}/_apis/connectionData"

        # Encode the personal access token in Base64
        token_bytes = f":{personal_access_token}".encode('ascii')
        base64_token = base64.b64encode(token_bytes).decode('ascii')

        # Set up the headers with the encoded token for authorization
        headers = {
            "Authorization": f"Basic {base64_token}"
        }

        # Make the GET request to obtain connection data
        response = requests.get(url, headers=headers)

        # Check if the response was successful
        if response.status_code == 200:
            connection_data = response.json()
            organization_name = connection_data.get('authenticatedUser', {}).get('properties', {}).get('Account', {}).get('$value')
            if organization_name:
                if 'ey' in organization_name.lower():
                    return True
                else:
                    return False
        else:
            return False

    except Exception as e:
        print(f"Ocorreu um erro ao obter o nome da organização: {str(e)}")
        return None

def get_all_relations(test_case_id):
    chave_decodificada = decodificar_chave()
    organization_url = 'https://dev.azure.com/BRPegaTeam/'
    personal_access_token = chave_decodificada

    credentials = BasicAuthentication('', personal_access_token)
    connection = Connection(base_url=organization_url, creds=credentials)
    wit_client = connection.clients.get_work_item_tracking_client()
        
    try:
        test_case = wit_client.get_work_item(test_case_id, expand='Relations')
        if test_case.relations is not None:
            for relation in test_case.relations:
                related_work_item_id = relation.url.split("/")[-1]
                type_rel = relation.rel.split('/')[-1]
                if type_rel == 'Microsoft.VSTS.Common.TestedBy-Reverse':
                    related_work_item = wit_client.get_work_item(related_work_item_id)
                    if related_work_item.fields['System.WorkItemType'] == 'Product Backlog Item':
                            # Imprimir o work item relacionado
                        related_title = related_work_item.fields['System.Title']
                        return related_title, test_case.fields['System.Title']
            response = messagebox.askyesno("Sem Relações", f"O caso de teste {test_case_id} não retornou nenhuma a US relacionada no DevOps. Deseja continuar?", icon='warning')
            if not response:
                return None
            else:
                if test_case.fields['System.Title'] != None:
                    return "", test_case.fields['System.Title']
                else:
                    return "", test_case_id
        else:
            response = messagebox.askyesno("Sem Relações", f"O caso de teste {test_case_id} não encontrou a US relacionada no DevOps. Deseja continuar?", icon='warning')
            if not response:
                return None
            else:
                print(test_case.fields['System.Title'])
                if test_case.fields['System.Title'] != None:
                    return "", test_case.fields['System.Title']
                else:
                    return "", test_case_id
    except:           
        response = messagebox.askyesno("Sem Relações", f"O caso de teste {test_case_id} não encontrou a US relacionada no DevOps. Deseja continuar?", icon='warning')
        if not response:
            return None
        else:
            return "", test_case_id
   
def ler_e_codificar_arquivo(novo_arquivo_path):
    try:
        # Abra o arquivo para leitura em modo binário
        with open(novo_arquivo_path, "rb") as arquivo:
            # Leia o conteúdo do arquivo
            conteudo = arquivo.read()

            # Codifique o conteúdo em base64
            arquivo_codificado_base64 = base64.b64encode(conteudo)

            # Retorne o arquivo codificado em base64 como uma string
            return arquivo_codificado_base64.decode("utf-8")

    except Exception as e:
        print("Ocorreu um erro ao ler e codificar o arquivo:", str(e))
        return None
    
# Função que cola as imagens em ordem no arquivo    
def criar_arquivos_com_imagens(diretorio_raiz, documento, id_personalizado):
    idx_paragrafo = next((idx for idx, p in enumerate(documento.paragraphs) if id_personalizado in p.text), None)
    arquivos = sorted(os.listdir(diretorio_raiz), key=lambda x: os.path.getmtime(os.path.join(diretorio_raiz, x)))

    first = 0
    if idx_paragrafo is not None and idx_paragrafo != len(documento.paragraphs) - 1:
        for nome_arquivo in arquivos:
            if nome_arquivo.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                if first == 0:
                    run = documento.paragraphs[idx_paragrafo + 1].add_run()
                else:
                    run = documento.paragraphs[idx_paragrafo + 2].add_run()
                caminho_imagem = os.path.join(diretorio_raiz, nome_arquivo)
                run.add_picture(caminho_imagem, width=Inches(9.0))
                first = 1

# Função que limpa as imagens do diretório selecionado
def limpar_imagens():
    diretorio_imagens = entries[1].get()
    if not diretorio_imagens:
        messagebox.showwarning("Campos Obrigatórios", "Por favor, preencha os campos Arquivo padrão e Diretórios de imagens.")
        return
    arquivos = os.listdir(diretorio_imagens)

    for arquivo in arquivos:
        caminho_arquivo = os.path.join(diretorio_imagens, arquivo)
        if arquivo.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            os.remove(caminho_arquivo)

    messagebox.showinfo("Limpeza Concluída", "Todos os arquivos de imagem foram removidos do diretório com sucesso.")

def on_checkbox_clicked():
    if checkbox_var.get():
        chave_decodificada = decodificar_chave()
        organization_url = 'https://dev.azure.com/BRPegaTeam/'
        personal_access_token = chave_decodificada
        validar_credenciais(organization_url, personal_access_token)
        return True
    else:
        return False

def criar_arquivo():
    try:
        documento_path = entries[0].get()
        diretorio_imagens = entries[1].get()
        
        if not documento_path or not diretorio_imagens:
            messagebox.showwarning("Campos Obrigatórios", "Por favor, preencha os campos Arquivo padrão e Diretórios de imagens.")
            return
        
        num_ct = entries[3].get()
        if not num_ct:
            messagebox.showwarning("Campo Obrigatório", "Por favor, preencha o campo ID do cenário de teste.")
            return
            
        imagens_encontradas = any(nome_arquivo.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')) for nome_arquivo in os.listdir(diretorio_imagens))
        
        
        
        documento = Document(documento_path)
        nome_tester = entries[2].get()
        num_ct = entries[3].get()
        ambiente = ambiente_dropdown.get()
        perfil = entries[5].get()
        bugs = entries[6].get()
        
        if on_checkbox_clicked():
            projeto_selecionado = project_dropdown.get()
            if projeto_selecionado:
                print("Projeto selecionado:", projeto_selecionado)
            else:
                messagebox.showwarning("Campo obrigatório", f"Por favor, preencha o campo projeto.")
                return
            if not decodificar_chave():
                messagebox.showwarning("Chave Não Cadastrada", "Chave do Azure DevOps não cadastrada.")
                return
            else:
                chave_decodificada = decodificar_chave()
                organization_url = 'https://dev.azure.com/BRPegaTeam/'
                personal_access_token = chave_decodificada
                if not validar_credenciais(organization_url, personal_access_token):
                    messagebox.showwarning("Credenciais inválidas", "Por favor, cadastre ou atualize sua chave do DevOps. Ela não é valida!")
                    return  
                else:
                    if imagens_encontradas:
                        titles = get_all_relations(num_ct)
                        if titles == None:
                            return
                    else:
                        resposta = messagebox.askquestion("Nenhuma Imagem Encontrada", "Nenhuma imagem foi encontrada no diretório especificado. Deseja continuar mesmo assim?")
                        if resposta == "no":
                            return
                        else:
                            titles = get_all_relations(num_ct)
                            if titles == None:
                                return
        else:
            titles = ["", num_ct]   
          
        criar_arquivos_com_imagens(diretorio_imagens, documento, "Evidências")
        
        for tabela in documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    text = celula.text
                    if '[Nome do Tester]' in text:
                        celula.text = ''
                        run = celula.paragraphs[0].add_run(text.replace('[Nome do Tester]', nome_tester))
                        run.font.name = 'EYInterstate Light'
                        run.font.size = Pt(16)
                    if '[Numero do CT]' in text:
                        celula.text = ''
                        if titles[1] != num_ct:
                            run = celula.paragraphs[0].add_run(text.replace('[Numero do CT]', num_ct + " - " + titles[1]))
                        else:
                            run = celula.paragraphs[0].add_run(text.replace('[Numero do CT]', num_ct))
                        run.font.name = 'EYInterstate Light'
                        run.font.size = Pt(14)
                    if '[US]' in text:
                        celula.text = ''
                        if titles[0] != "":
                            run = celula.paragraphs[0].add_run(text.replace('[US]', titles[0]))
                        else:
                            run = celula.paragraphs[0].add_run(text.replace('[US]', ''))
                        run.font.name = 'EYInterstate Light'
                        run.font.size = Pt(14)
                    if '[Ambiente]' in text:
                        celula.text = ''
                        run = celula.paragraphs[0].add_run(text.replace('[Ambiente]', ambiente))
                        run.font.name = 'EYInterstate Light'
                        run.font.size = Pt(16)
                    if '[Perfil]' in text:
                        celula.text = ''
                        run = celula.paragraphs[0].add_run(text.replace('[Perfil]', perfil))
                        run.font.name = 'EYInterstate Light'
                        run.font.size = Pt(16)
                    if '[Bugs]' in text:
                        celula.text = ''
                        run = celula.paragraphs[0].add_run(text.replace('[Bugs]', bugs))
                        run.font.name = 'EYInterstate Light'
                        run.font.size = Pt(16)                
                    if '[Resultado]' in text:
                        celula.text = ''
                        if len(bugs) == 0:
                            run = celula.paragraphs[0].add_run(text.replace('[Resultado]', '☒Passed  ☐Failed'))
                        else:
                            run = celula.paragraphs[0].add_run(text.replace('[Resultado]', '☐Passed  ☒Failed'))
                        run.font.name = 'EYInterstate Light'
                        run.font.size = Pt(16)

        # Obtendo o nome base do documento (sem o diretório)
        base_nome = os.path.basename(documento_path)

        # Dividindo o nome do arquivo em nome e extensão
        nome, extensao = os.path.splitext(base_nome)

        # Criando o novo nome do arquivo, substituindo o nome original pelo desejado
        novo_nome = f"{nome}_{num_ct}.docx" 
        novo_arquivo_path = os.path.join(diretorio_imagens, novo_nome)

        documento.save(novo_arquivo_path)
                        # Exemplo de uso
        try:
            import shutil
            if on_checkbox_clicked():            
                
                # Copiar o arquivo para o diretório atual
                diretorio_pai = os.getcwd()
                # Caminho completo do destino do arquivo (diretório pai)
                caminho_arquivo_destino = os.path.join(diretorio_pai, novo_nome)
                
                caminho_arquivo_origem = os.path.join(os.getcwd(), novo_arquivo_path)
                shutil.copy(caminho_arquivo_origem, caminho_arquivo_destino)
                
                # Copiando o arquivo
                upload_url = f"https://dev.azure.com/BRPegaTeam/{projeto_selecionado}/_apis/wit/attachments?fileName={novo_nome}&api-version=6.0"
                attachment_url = upload_attachment(novo_nome, upload_url)
                     # Remover o arquivo temporário
                os.remove(novo_nome)
                if attachment_url != None:
                    work_item_url = f"https://dev.azure.com/BRPegaTeam/{projeto_selecionado}/_apis/wit/workitems/{num_ct}?api-version=6.1-preview.3"
                    add_attachment_to_work_item(num_ct, attachment_url, work_item_url)
                    if add_attachment_to_work_item != None:
                        messagebox.showinfo("Sucesso", f"Documento criado com sucesso! Salvo em: {novo_arquivo_path}.\n\n Evidência anexada com sucesso ao caso de teste.")
                    else:
                        messagebox.showinfo("Sucesso", f"Documento criado com sucesso! Salvo em: {novo_arquivo_path}.")
                        messagebox.showerror("Erro", f"Ocorreu um erro ao anexar a evidência no caso de teste")
                else:
                    messagebox.showinfo("Sucesso", f"Documento criado com sucesso! Salvo em: {novo_arquivo_path}.")
                    messagebox.showerror("Erro", f"Ocorreu um erro ao anexar a evidência no caso de teste")
            else:
                messagebox.showinfo("Sucesso", f"Documento criado com sucesso! Salvo em: {novo_arquivo_path}.")

        except Exception as e:
            messagebox.showinfo("Sucesso", f"Documento criado com sucesso!\nSalvo em: {novo_arquivo_path}")
            messagebox.showerror("Erro", f"Ocorreu um erro ao anexar a evidência no caso de teste: {str(e)}")
        


        # Show a success message

    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao gerar o documento:\n{str(e)}")
        print("Erro:", str(e))


def upload_attachment(filepath, url):
    """Faz upload de um arquivo como anexo para o Azure DevOps."""
    # Lendo o arquivo
    with open(filepath, 'rb') as f:
        file_content = f.read()
    
    print(file_content)
    # Encoder no formato username:PAT
    user_pat = f"{''}:{decodificar_chave()}"
    b64_user_pat = base64.b64encode(user_pat.encode()).decode()
    headers = {
    'Accept': 'application/json-patch+json',
    'Content-Type': 'application/octet-stream',
    'Authorization': f'Basic {b64_user_pat}'
    }
        
    # Fazendo a requisição de POST
    r = requests.post(url, headers=headers, data=file_content)
    # Parsing da resposta recebida
    try:
        response = r.json()
    except ValueError:  
        messagebox.showerror("Erro", f"Ocorreu um erro ao anexar a evidência no caso de teste: {ValueError}")
        print("Erro no parsing da resposta")
        return None
    # Verifica o código de status HTTP
    if r.status_code != 201:
        messagebox.showerror("Erro", f"Ocorreu um erro ao anexar a evidência no caso de teste, a verificação não retornou código 201.")
        print(f"Erro ao fazer upload do anexo: {response}")
        return None
    else:
        print(f"Anexo carregado com sucesso: {response['url']}")
        return response['url']
import json
def add_attachment_to_work_item(work_item_id, attachment_url, url):
    # Construindo url
    url = f"https://dev.azure.com/BRPegaTeam/Gestão%20de%20OPME/_apis/wit/workitems/{work_item_id}?api-version=6.1-preview.3"
    user_pat = f"{''}:{decodificar_chave()}"
    b64_user_pat = base64.b64encode(user_pat.encode()).decode()
    headers = {
        'Content-Type': 'application/json-patch+json',
        'Authorization': f'Basic {b64_user_pat}'
    }

    # Preparando o corpo da requisição
    relation = {
        "op": "add",
        "path": "/relations/-",
        "value": {
            "rel": "AttachedFile",
            "url": attachment_url,
             "attributes": {
            "comment": "Evidência de teste"
        }            
        }
    }

    body = json.dumps([relation])

    # Envia a requisição PATCH para a API
    r = requests.patch(url, headers=headers, data=body)

    # Verifica se a requisição foi bem-sucedida
    if r.status_code != 200:
        print(f"Erro ao adicionar anexo ao Work Item: {r.text}")
        return None
    else:
        print("Anexo adicionado com sucesso ao Work Item.")

        
janela = ThemedTk(theme="arc")
janela.geometry("565x455")
janela.iconbitmap("ey.ico")
janela.title("Gerador de KKKKKKKKKK")

style = ttk.Style()
style.configure("TEntry", padding=(5, 4, 5, 4), relief="flat")
style.configure("TCombobox", padding=(5, 4, 5, 4), relief="flat")
style.configure("TButton", padding=(10, 5, 10, 5), relief="flat")

labels_text = ["Arquivo padrão", "Diretório de Imagens", "Nome do Tester", "ID do cenário de teste", "Ambiente", "Produto", "Perfil", "Bugs"]
entries = [ttk.Entry(janela) for _ in labels_text]

for i in range(len(labels_text)):
    ttk.Label(janela, text=labels_text[i]).grid(row=i + 1, column=0, padx=20, pady=10, sticky='w')
    if labels_text[i] == "Ambiente" or labels_text[i] == "Produto":
        i+1
    else:
        entries[i].grid(row=i + 1, column=1, padx=2, pady=10)

btn_browse = ttk.Button(janela, text="Procurar", command=on_browse)
btn_browse.grid(row=1, column=2, padx=0, pady=10, sticky='w')

btn_select_image_dir = ttk.Button(janela, text="Procurar", command=on_select_image_dir)
btn_select_image_dir.grid(row=2, column=2, padx=0, pady=0, sticky='w')

btn_clear_images = ttk.Button(janela, text="Limpar imagens", command=limpar_imagens)
btn_clear_images.grid(row=2, column=3, padx=5, pady=10, sticky='w')
         
rbtns_text = ["TU", "TI", "HML"]
amb_var = tk.StringVar(value='TI')
ambiente_dropdown = ttk.Combobox(janela, textvariable=amb_var, values=rbtns_text, style="TCombobox")
ambiente_dropdown.grid(row=5, column=1, padx=20, pady=10, sticky='e')
ambiente_dropdown.configure(width=18)

ambiente_dropdown.bind("<<ComboboxSelected>>", desativar_selecao)

btn_ativar = ttk.Button(janela, text="Gerar", command=criar_arquivo)
btn_ativar.grid(row=9, column=1, padx=20, pady=10)
btn_ativar.configure(width=20) 

# Carregar a imagem
image = Image.open("chave.png")

# Redimensionar a imagem para o tamanho desejado
width, height = 14, 14
image = image.resize((width, height))  # Correção aqui

# Converter a imagem para um objeto PhotoImage
photo = ImageTk.PhotoImage(image)

# Botão para cadastrar/renovar chave do Azure DevOps
btn_chave_devops = ttk.Button(janela, command=cadastrar_renovar_chave, image=photo, compound="left")
btn_chave_devops.grid(row=3, column=3, padx=65, pady=0)
btn_chave_devops.configure(width=1)


# Checkbox para integração com Azure DevOps
checkbox_var = tk.BooleanVar(value=True)
checkbox = ttk.Checkbutton(janela, text="Integrar com DevOps", variable=checkbox_var, command=on_checkbox_clicked)
checkbox.grid(row=3, column=2, columnspan=2, padx=(0, 20), pady=10, sticky='w')  # Ocupa duas colunas


# Função para obter os projetos do Azure DevOps
def get_devops_projects():
    if on_checkbox_clicked():
            if not decodificar_chave():
                messagebox.showwarning("Chave Não Cadastrada", "Chave do Azure DevOps não cadastrada.")
                return
            else:
                chave_decodificada = decodificar_chave()
                organization_url = 'https://dev.azure.com/BRPegaTeam/'
                personal_access_token = chave_decodificada
                if not validar_credenciais(organization_url, personal_access_token):
                    messagebox.showwarning("Credenciais inválidas", "Por favor, cadastre ou atualize sua chave do DevOps. Ela não é valida!")
                    return  
                else:
                    try:
                        # URL para obter os projetos
                        url = "https://dev.azure.com/BRPEGAteam/_apis/projects?api-version=6.1-preview.4"
                        
                        
                        # Encode the personal access token in Base64
                        token_bytes = f":{decodificar_chave()}".encode('ascii')
                        base64_token = base64.b64encode(token_bytes).decode('ascii')

                        # Set up the headers with the encoded token for authorization
                        headers = {
                            "Authorization": f"Basic {base64_token}"
                        }
                        # Fazendo a requisição GET
                        response = requests.get(url, headers=headers)
                        # Verificando se a resposta foi bem-sucedida
                        if response.status_code == 200:
                            # Obtendo os projetos da resposta JSON
                            projects = response.json().get("value", [])
                            # Retornando apenas os nomes dos projetos
                            return [project['name'] for project in projects]
                        else:
                            print("Erro ao obter os projetos:", response.status_code)
                            return []
                    except Exception as e:
                        print("Erro ao obter os projetos:", str(e))
                        return []

# Função para atualizar os valores do dropdown com os projetos do Azure DevOps
def update_project_dropdown():
    # Verificando se o checkbox está marcado
    if checkbox_var.get():
        # Obtendo os projetos do Azure DevOps
        projects = get_devops_projects()
        # Atualizando os valores do dropdown
        project_dropdown.config(values=projects)
    else:
        # Se o checkbox não estiver marcado, definir os valores do dropdown como vazio
        project_dropdown.config(values=[])


# Criando um checkbox para integração com Azure DevOps
checkbox_var = tk.BooleanVar(value=False)  # Inicialmente não marcado
checkbox = ttk.Checkbutton(janela, text="Integrar com DevOps", variable=checkbox_var, command=update_project_dropdown)
checkbox.grid(row=3, column=2, columnspan=2, padx=(0, 20), pady=10, sticky='w')  # Ocupa duas colunas

# Criando um dropdown para selecionar o projeto do Azure DevOps
project_dropdown = ttk.Combobox(janela, state="readonly")
project_dropdown.grid(row=6, column=1, padx=20, pady=10, sticky='e')
project_dropdown.configure(width=18)


project_dropdown.bind("<<ComboboxSelected>>", desativar_selecao)


janela.mainloop()

